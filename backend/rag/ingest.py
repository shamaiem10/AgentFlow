import pymupdf as fitz
import pdfplumber
from docx import Document as DocxDocument
import openpyxl
import csv
import os
import easyocr
from PIL import Image
import io
from agents.image_captioning import caption_image_bytes

ocr_reader = easyocr.Reader(['en'])

def table_to_text(table):
    """Converts a pdfplumber table (list of rows) into readable | separated text."""
    lines = []
    for row in table:
        cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
        lines.append(" | ".join(cleaned_row))
    return "\n".join(lines)

def extract_tables_from_pdf(file_path):
    """Returns a dict: {page_number: [table_text, table_text, ...]}"""
    tables_by_page = {}
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if tables:
                tables_by_page[page_num] = [table_to_text(t) for t in tables]
    return tables_by_page

def parse_pdf(file_path):
    doc = fitz.open(file_path)
    full_text = ""

    # extract tables separately first, keyed by page number
    tables_by_page = extract_tables_from_pdf(file_path)

    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            full_text += text + "\n"

        # insert any tables found on this page
        if page_num in tables_by_page:
            for i, table_text in enumerate(tables_by_page[page_num]):
                full_text += f"\n[Table {i+1} on page {page_num+1}]:\n{table_text}\n"

        # extract and caption any embedded images on this page
        image_list = page.get_images(full=True)
        if image_list:
            print(f"Page {page_num + 1}: found {len(image_list)} image(s), captioning...")

        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            try:
                print(f"  Captioning image {img_index + 1}/{len(image_list)}...")
                caption = caption_image_bytes(image_bytes)
                full_text += f"\n[Image description: {caption}]\n"
                print(f"  Done.")
            except Exception as e:
                print(f"  Warning: failed to caption an image — {e}")

    if len(full_text.strip()) < 50:
        full_text = ocr_pdf(file_path)

    doc.close()
    return full_text

def ocr_pdf(file_path):
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_bytes))
        results = ocr_reader.readtext(image, detail=0)
        full_text += " ".join(results) + "\n"
    doc.close()
    return full_text

def parse_image(file_path):
    results = ocr_reader.readtext(file_path, detail=0)
    return " ".join(results)

def parse_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            full_text += row_text + "\n"
    return full_text

def parse_xlsx(file_path):
    workbook = openpyxl.load_workbook(file_path, data_only=True)
    full_text = ""
    for sheet in workbook.worksheets:
        full_text += f"\n--- Sheet: {sheet.title} ---\n"
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell) for cell in row if cell is not None]
            if row_values:
                full_text += " | ".join(row_values) + "\n"
    return full_text

def parse_csv(file_path):
    full_text = ""
    with open(file_path, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            full_text += " | ".join(row) + "\n"
    return full_text

def parse_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return parse_xlsx(file_path)
    elif ext == ".csv":
        return parse_csv(file_path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        return parse_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")